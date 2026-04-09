import streamlit as st
import pandas as pd
from openai import OpenAI
import os
from fpdf import FPDF

# --- 1. CẤU HÌNH ---
st.set_page_config(page_title="Hệ thống Thủy lợi Sơn La", layout="wide")

# --- 2. HÀM TẠO PDF (BẢN FIX LỖI CHIỀU NGANG) ---
import streamlit as st
import pandas as pd
from openai import OpenAI
import os
from docx import Document
from io import BytesIO

# --- 1. CẤU HÌNH ---
st.set_page_config(page_title="Hệ thống Thủy lợi Sơn La", layout="wide")

# --- 2. HÀM TẠO FILE WORD (DỄ SỬA, KHÔNG LỖI FONT) ---
def tao_file_word(tra_loi, cau_hoi, ten_ct):
    doc = Document()
    
    # Tiêu đề quốc hiệu
    p = doc.add_paragraph()
    p.alignment = 1 # Căn giữa
    run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n")
    run.bold = True
    
    doc.add_heading("BIÊN BẢN TRA CỨU NGHIỆP VỤ THỦY LỢI", level=1).alignment = 1
    
    doc.add_paragraph(f"Tên công trình: {ten_ct}")
    doc.add_paragraph(f"Nội dung câu hỏi: {cau_hoi}")
    doc.add_paragraph("-" * 20)
    
    # Nội dung AI trả lời
    doc.add_heading("KẾT QUẢ TRA CỨU (Cần rà soát lại):", level=2)
    doc.add_paragraph(tra_loi)
    
    doc.add_paragraph("\n" + "_" * 30)
    doc.add_paragraph("Đại diện Chi nhánh Thủy lợi số 5: ....................................")
    doc.add_paragraph("Đại diện UBND phường/bản: ........................................")
    doc.add_paragraph("Cán bộ địa bàn: .................................................")
    doc.add_paragraph("Hộ gia đình vi phạm: .............................................")
    
    p_date = doc.add_paragraph(f"\nNgày lập biên bản: {pd.Timestamp.now().strftime('%d/%m/%Y')}")
    p_date.alignment = 2 # Căn phải
    
    # Lưu vào bộ nhớ đệm để Streamlit tải về
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 3. HÀM NẠP DỮ LIỆU ---
@st.cache_resource
def nap_du_lieu():
    context = ""
    if os.path.exists("data"):
        from pypdf import PdfReader
        for file in os.listdir("data"):
            if file.endswith(".pdf"):
                try:
                    reader = PdfReader(os.path.join("data", file))
                    for page in reader.pages: context += page.extract_text() + "\n"
                except: continue
    return context

# --- 4. HỆ THỐNG ---
if "OPENAI_API_KEY" in st.secrets:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    context_data = nap_du_lieu()
else:
    st.error("Chưa có API Key!")
    st.stop()

# --- 5. GIAO DIỆN ---
st.title("🌊 Trợ lý AI Thủy Lợi (Xuất file Word)")
col1, col2 = st.columns([1, 1])

with col1:
    st.header("📍 Thông số & Bản đồ")
    try:
        file_excel = [f for f in os.listdir("specs") if f.endswith(".xlsx")][0]
        df = pd.read_excel(os.path.join("specs", file_excel))
        df.columns = df.columns.str.strip()
        col_ten = df.select_dtypes(include=['object']).columns[0]
        ten_ct = st.selectbox("Chọn công trình:", df[col_ten].unique())
        row = df[df[col_ten] == ten_ct].iloc[0]
        for c in df.columns:
            if c.lower() not in ['lat', 'vĩ', 'lon', 'kinh']:
                st.write(f"**{c}:** {row[c]}")
        map_url = f"https://www.google.com/maps?q={row.get('lat', 21.3)},{row.get('lon', 103.9)}&output=embed&t=k"
        st.components.v1.iframe(map_url, height=400)
    except: st.warning("Đang tải dữ liệu...")

with col2:
    st.header("💬 Hỏi đáp & Trích xuất")
    hoi = st.text_input("Nhập câu hỏi để AI soạn thảo:")
    if hoi:
        with st.spinner("AI đang soạn thảo biên bản..."):
            res = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": f"Dữ liệu: {context_data[:3000]}"},
                    {"role": "user", "content": hoi}
                ]
            )
            tra_loi = res.choices[0].message.content
            st.write(tra_loi)
            
            # NÚT TẢI FILE WORD
            st.markdown("---")
            word_data = tao_file_word(tra_loi, hoi, ten_ct)
            st.download_button(
                label="📥 Tải Biên bản (File Word để sửa)",
                data=word_data,
                file_name=f"Bien_ban_{ten_ct}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# --- 3. HÀM NẠP DỮ LIỆU (FIX LỖI NAMEERROR) ---
@st.cache_resource
def nap_du_lieu_van_ban():
    context = ""
    if os.path.exists("data"):
        from pypdf import PdfReader
        for file in os.listdir("data"):
            if file.endswith(".pdf"):
                try:
                    reader = PdfReader(os.path.join("data", file))
                    for page in reader.pages:
                        context += page.extract_text() + "\n"
                except: continue
    return context

# --- 4. KẾT NỐI HỆ THỐNG ---
if "OPENAI_API_KEY" in st.secrets:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    if "context" not in st.session_state:
        st.session_state.context = nap_du_lieu_van_ban()
else:
    st.error("Chưa có API Key!")
    st.stop()

# --- 5. GIAO DIỆN ---
st.title("🌊 Trợ lý AI Ngành Thủy Lợi (Sơn La)")
col1, col2 = st.columns([1, 1])

with col1:
    st.header("📍 Bản đồ & Thông số")
    try:
        file_excel = [f for f in os.listdir("specs") if f.endswith(".xlsx")][0]
        df = pd.read_excel(os.path.join("specs", file_excel))
        df.columns = df.columns.str.strip()

        col_ten = df.select_dtypes(include=['object']).columns[0]
        col_lat = [c for c in df.columns if 'lat' in c.lower() or 'vĩ' in c.lower()][0]
        col_lon = [c for c in df.columns if 'lon' in c.lower() or 'kinh' in c.lower()][0]

        ten_ct = st.selectbox("Chọn công trình:", df[col_ten].unique())
        row = df[df[col_ten] == ten_ct].iloc[0]
        
        # Hiện mọi thông số trong Excel
        for c in df.columns:
            if c not in [col_lat, col_lon]:
                st.write(f"**{c}:** {row[c]}")
        
        map_url = f"https://www.google.com/maps?q={row[col_lat]},{row[col_lon]}&output=embed&t=k"
        st.components.v1.iframe(map_url, height=400)
    except: st.warning("Đang tải dữ liệu...")

with col2:
    st.header("💬 Hỏi đáp Trợ lý AI")
    hoi = st.text_input("Nhập câu hỏi:")
    if hoi:
        with st.spinner("AI đang trả lời..."):
            res = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": f"Dữ liệu: {st.session_state.context[:3000]}"},
                    {"role": "user", "content": hoi}
                ]
            )
            tra_loi = res.choices[0].message.content
            st.write(tra_loi)
            
            # HIỆN NÚT TẢI PDF
            st.markdown("---")
            pdf_data = tao_pdf_bien_ban(tra_loi, hoi, ten_ct)
            st.download_button("📥 Tải Biên bản PDF", data=pdf_data, file_name="Bien_ban.pdf", mime="application/pdf")